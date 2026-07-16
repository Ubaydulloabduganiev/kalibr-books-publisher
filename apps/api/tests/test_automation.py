"""Tests for the content-plan automation pipeline."""

from __future__ import annotations

import io
from datetime import datetime, timezone
from pathlib import Path

import pytest

from kalibr_publisher.core import store
from kalibr_publisher.services import automation
from kalibr_publisher.services.automation import run_automation
from kalibr_publisher.services.document_parser import (
    PlanItem,
    ParsedPlan,
    parse_plan,
    parse_text,
)


class _FakeImage:
    def __init__(self, data: bytes = b"fakeimg", mime: str = "image/png"):
        self.data = data
        self.mime = mime

    def generate(self, prompt: str, max_retries: int = 3):  # noqa: D401
        return _FakeImage()


class _FakeCaption:
    def rewrite_caption(self, text: str, language: str = "uz"):
        from kalibr_publisher.integrations.gemini import CaptionResult

        return CaptionResult(text=f"[uz] {text}", parse_mode="HTML")


class _FakeSettings:
    telegram_default_channel = "@kalibr_books"
    media_root = "storage_test/media"

    def __init__(self):
        Path(self.media_root).mkdir(parents=True, exist_ok=True)


def test_parse_text_splits_lines():
    plan = parse_text("- Learn phrasal verbs\n* Daily vocabulary\n3) Grammar tips")
    assert [i.title for i in plan.items] == [
        "Learn phrasal verbs",
        "Daily vocabulary",
        "Grammar tips",
    ]


def _make_pdf(text: str) -> bytes:
    lines = text.split("\n")
    y = 700
    ops = [f"BT /F1 12 Tf 50 {y} Td ({ln}) Tj ET" for y, ln in ((yy, l) for yy, l in zip(range(700, 0, -20), lines))]
    stream = " ".join(ops)
    objs = [
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj",
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj",
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 800 800]/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>endobj",
        f"4 0 obj<</Length {len(stream)}>>stream\n{stream}\nendstream endobj".encode(),
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj",
    ]
    pdf = b""
    offsets = []
    for o in objs:
        offsets.append(len(pdf))
        pdf += o + b"\n"
    xref = b"xref\n0 6\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n".encode()
    xref += b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n" + str(offsets[0]).encode() + b"\n%%EOF"
    return b"%PDF-1.4\n" + pdf + xref


def test_parse_docx_and_pdf():
    import docx
    from docx import Document as RealDoc

    # real docx with an empty paragraph (exercises the `if not title: continue` branch)
    d = RealDoc()
    d.add_paragraph("Topic A")
    d.add_paragraph("")  # empty -> skipped
    d.add_paragraph("Topic B")
    buf = io.BytesIO()
    d.save(buf)
    plan = parse_plan("plan.docx", buf.getvalue())
    assert [i.title for i in plan.items] == ["Topic A", "Topic B"]

    # real pdf (exercises parse_pdf real path)
    plan = parse_plan("plan.pdf", _make_pdf("Line 1\nLine 2"))
    assert "Line 1" in [i.title for i in plan.items]

    # txt path
    plan = parse_plan("plan.txt", "Plain one\nPlain two".encode())
    assert [i.title for i in plan.items] == ["Plain one", "Plain two"]


def test_run_automation_creates_scheduled_posts(monkeypatch, tmp_path):
    settings = _FakeSettings()
    settings.media_root = str(tmp_path / "media")
    monkeypatch.setattr(automation, "GeminiClient", lambda *a, **k: _FakeCaption())
    monkeypatch.setattr(automation, "GeminiImageClient", lambda *a, **k: _FakeImage())

    plan = ParsedPlan(items=[PlanItem(title="Vocab 1"), PlanItem(title="Vocab 2")])
    created = run_automation(plan, settings=settings, language="uz", stagger_hours=24)

    assert len(created) == 2
    # staggered one day apart
    t0 = datetime.fromisoformat(created[0]["run_at"])
    t1 = datetime.fromisoformat(created[1]["run_at"])
    assert (t1 - t0).total_seconds() == 24 * 3600
    # persisted to store
    assert len(store.list_posts()) >= 2
    # image file written
    assert Path(created[0]["image"]).exists()


def test_run_automation_caption_fallback(monkeypatch, tmp_path):
    class _EmptyCaption:
        def rewrite_caption(self, text: str, language: str = "uz"):
            from kalibr_publisher.integrations.gemini import CaptionResult

            return CaptionResult(text="", parse_mode="HTML")  # empty -> fallback

    settings = _FakeSettings()
    settings.media_root = str(tmp_path / "media")
    monkeypatch.setattr(automation, "GeminiClient", lambda *a, **k: _EmptyCaption())
    monkeypatch.setattr(automation, "GeminiImageClient", lambda *a, **k: _FakeImage())

    plan = ParsedPlan(items=[PlanItem(title="Fallback Topic")])
    created = run_automation(plan, settings=settings, language="uz")
    assert created[0]["text"] == "Fallback Topic"
